from __future__ import annotations

import io
import re
import unicodedata
from pathlib import Path

from docx import Document

ALLOWED_EXTENSIONS = {".pdf", ".docx"}
MAX_RESUME_BYTES = 5 * 1024 * 1024

# Icon-font / Private Use Area glyphs often appear as boxes or CJK-looking junk.
_PRIVATE_USE = re.compile(r"[\uE000-\uF8FF]")
_CONTROL_CHARS = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]")


class ResumeParseError(ValueError):
    """Raised when a resume cannot be validated or parsed."""


def _extension(filename: str) -> str:
    return Path(filename).suffix.lower()


def validate_resume_filename(filename: str | None) -> str:
    if not filename or not filename.strip():
        raise ResumeParseError("Resume filename is required")

    ext = _extension(filename)
    if ext not in ALLOWED_EXTENSIONS:
        raise ResumeParseError("Resume must be a PDF or DOCX file")
    return ext


def parse_resume_bytes(data: bytes, filename: str) -> str:
    if not data:
        raise ResumeParseError("Resume file is empty")
    if len(data) > MAX_RESUME_BYTES:
        raise ResumeParseError("Resume must be 5 MB or smaller")

    ext = validate_resume_filename(filename)

    if ext == ".pdf":
        text = _parse_pdf(data)
    else:
        text = _parse_docx(data)

    cleaned = _normalize_text(text)
    if not cleaned:
        raise ResumeParseError("Could not extract text from the resume")
    return cleaned


def _parse_pdf(data: bytes) -> str:
    # PyMuPDF handles custom resume fonts much more reliably than pypdf.
    try:
        import fitz  # pymupdf
    except ImportError:
        fitz = None

    if fitz is not None:
        try:
            with fitz.open(stream=data, filetype="pdf") as doc:
                parts = [page.get_text("text") for page in doc]
            text = "\n".join(parts)
            if _normalize_text(text):
                return text
        except Exception:  # noqa: BLE001
            pass

    return _parse_pdf_pypdf(data)


def _parse_pdf_pypdf(data: bytes) -> str:
    from pypdf import PdfReader

    try:
        reader = PdfReader(io.BytesIO(data))
    except Exception as exc:  # noqa: BLE001 - library raises varied errors
        raise ResumeParseError("Could not read PDF resume") from exc

    parts: list[str] = []
    for page in reader.pages:
        try:
            page_text = page.extract_text() or ""
        except Exception:  # noqa: BLE001
            page_text = ""
        if page_text:
            parts.append(page_text)
    return "\n".join(parts)


def _parse_docx(data: bytes) -> str:
    try:
        document = Document(io.BytesIO(data))
    except Exception as exc:  # noqa: BLE001
        raise ResumeParseError("Could not read DOCX resume") from exc

    parts: list[str] = []
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def _normalize_text(text: str) -> str:
    # Null bytes make editors treat the file as binary/UTF-16 → Chinese-looking junk.
    text = text.replace("\x00", "")
    text = unicodedata.normalize("NFKC", text)
    text = text.replace("\u00a0", " ").replace("\u202f", " ").replace("\ufeff", "")
    text = _PRIVATE_USE.sub(" ", text)
    text = _CONTROL_CHARS.sub("", text)
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+\n", "\n", text)
    text = re.sub(r"[ \t]{2,}", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def safe_stem(filename: str) -> str:
    stem = Path(filename).stem
    cleaned = re.sub(r"[^A-Za-z0-9._-]+", "_", stem).strip("._")
    return cleaned[:80] or "resume"
